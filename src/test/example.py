from docx import Document

import math2docx

document = Document()
p = document.add_paragraph()

#LaTex Math Formula
latex_ = r"\cos (2\theta) = \cos^2 \theta - \sin^2 \theta"

math2docx.add_math(p, latex_ )

p2 = document.add_paragraph()
math2docx.add_math(p2, latex_, False)
document.save("new_file.docx")
